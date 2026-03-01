"""Document loading and text chunking (PDF/DOCX/TXT)."""
from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader


def load_pdf(file_path: str | Path) -> list[Document]:
    """Extract text from every page of a PDF and return as Document list."""
    reader = PdfReader(str(file_path))
    docs: list[Document] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": Path(file_path).name, "page": i + 1},
                )
            )
    return docs


def load_docx(file_path: str | Path) -> list[Document]:
    """Extract text from DOCX and return as a single Document."""
    docx = DocxDocument(str(file_path))
    text = "\n".join(p.text for p in docx.paragraphs if p.text and p.text.strip())
    if not text.strip():
        return []
    return [
        Document(
            page_content=text,
            metadata={"source": Path(file_path).name, "page": None},
        )
    ]


def load_txt(file_path: str | Path) -> list[Document]:
    """Extract text from TXT using common encodings and return as Document list."""
    text = ""
    for encoding in ("utf-8", "utf-8-sig", "cp1252"):
        try:
            text = Path(file_path).read_text(encoding=encoding)
            break
        except UnicodeDecodeError:
            continue
    if not text.strip():
        return []
    return [
        Document(
            page_content=text,
            metadata={"source": Path(file_path).name, "page": None},
        )
    ]


def split_documents(
    docs: list[Document],
    chunk_size: int = 512,
    chunk_overlap: int = 80,
) -> list[Document]:
    """Split documents into overlapping chunks for retrieval.

    Uses RecursiveCharacterTextSplitter which respects paragraph → sentence →
    word boundaries — good for preserving n-gram context used by BM25.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(docs)


def process_pdf(
    file_path: str | Path,
    chunk_size: int = 512,
    chunk_overlap: int = 80,
) -> list[Document]:
    """Full pipeline: load PDF → split into chunks."""
    docs = load_pdf(file_path)
    return split_documents(docs, chunk_size=chunk_size, chunk_overlap=chunk_overlap)


def process_document(
    file_path: str | Path,
    chunk_size: int = 512,
    chunk_overlap: int = 80,
) -> list[Document]:
    """Full pipeline for supported file types: PDF, DOCX, TXT."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        docs = load_pdf(file_path)
    elif ext == ".docx":
        docs = load_docx(file_path)
    elif ext == ".txt":
        docs = load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return split_documents(docs, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
